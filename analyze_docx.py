#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to analyze and extract bilingual content from DOCX file
"""
import docx
import json

def analyze_docx(filename):
    doc = docx.Document(filename)

    print(f"=== Document Analysis ===")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    print()

    # Show first few paragraphs
    print("=== First 10 paragraphs ===")
    for i, para in enumerate(doc.paragraphs[:10]):
        if para.text.strip():
            print(f"{i}: {para.text[:100]}")
    print()

    # Analyze tables structure
    if doc.tables:
        print("=== Tables Structure ===")
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx}:")
            print(f"  Rows: {len(table.rows)}")
            print(f"  Columns: {len(table.rows[0].cells) if table.rows else 0}")

            # Show first few rows
            print(f"  First 3 rows:")
            for row_idx, row in enumerate(table.rows[:3]):
                cells_text = [cell.text[:50] for cell in row.cells]
                print(f"    Row {row_idx}: {cells_text}")

    return doc

if __name__ == "__main__":
    doc = analyze_docx("LB17-11 viss tulkojums.docx")
