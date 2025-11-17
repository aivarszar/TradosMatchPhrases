#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Create DOCX output file with corrected segments
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_corrected_docx():
    """Create DOCX with corrected segments"""

    # Load corrections
    with open('translation_corrections.json', 'r', encoding='utf-8') as f:
        corrections = json.load(f)

    # Create new document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add title
    title = doc.add_heading('Tulkojuma Labojumi / Translation Corrections', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add document info
    info_para = doc.add_paragraph()
    info_para.add_run('Dokuments: ').bold = True
    info_para.add_run('UM_6789320-1EN_r1_for_RevolutionVibe_Apex5.0-1107\n')
    info_para.add_run('Avota fails: ').bold = True
    info_para.add_run('LB17-11 viss tulkojums.docx\n')
    info_para.add_run('Labojumu skaits: ').bold = True
    info_para.add_run(f'{len(corrections)}\n')
    info_para.add_run('Datums: ').bold = True
    info_para.add_run(__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Spacer

    # Add summary
    summary_heading = doc.add_heading('Kopsavilkums / Summary', 1)

    summary = doc.add_paragraph()
    summary.add_run('• Kopā analizēti: ').bold = True
    summary.add_run('2662 segmenti\n')
    summary.add_run('• Identificētas problēmas: ').bold = True
    summary.add_run('30 segmentos (1.1%)\n')
    summary.add_run('• Veikti labojumi: ').bold = True
    summary.add_run('8 segmentos (0.3%)\n')
    summary.add_run('• Kvalitātes līmenis: ').bold = True
    summary.add_run('99.7%')

    doc.add_paragraph()  # Spacer

    # Add corrections table
    corrections_heading = doc.add_heading('Detalizēti Labojumi', 1)

    # Sort corrections by segment ID
    corrections.sort(key=lambda x: x['segment_id'])

    # Create table for each correction
    for i, corr in enumerate(corrections, 1):
        # Add correction header
        header = doc.add_heading(f'Labojums #{i} - Segments #{corr["segment_id"]}', 2)

        # Add severity and issue type
        info = doc.add_paragraph()
        info.add_run('Prioritāte: ').bold = True
        severity_text = {
            'HIGH': 'Augsta',
            'MEDIUM': 'Vidēja',
            'LOW': 'Zema'
        }.get(corr['severity'], corr['severity'])
        severity_run = info.add_run(f'{severity_text}    ')
        if corr['severity'] == 'HIGH':
            severity_run.font.color.rgb = RGBColor(231, 76, 60)
        elif corr['severity'] == 'MEDIUM':
            severity_run.font.color.rgb = RGBColor(243, 156, 18)
        else:
            severity_run.font.color.rgb = RGBColor(52, 152, 219)

        info.add_run('Problēmas veids: ').bold = True
        info.add_run(corr['issue_type'])

        # Create table for this correction
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.2)
            row.cells[1].width = Inches(5.3)

        # Row 1: Source
        cell = table.rows[0].cells[0]
        cell.text = 'EN (Avotteksts)'
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, '3498DB')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        table.rows[0].cells[1].text = corr['original_source']

        # Row 2: Original translation
        cell = table.rows[1].cells[0]
        cell.text = 'LV (Oriģināls)'
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'E74C3C')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        table.rows[1].cells[1].text = corr['original_target']

        # Row 3: Corrected translation
        cell = table.rows[2].cells[0]
        cell.text = 'LV (Labots)'
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, '2ECC71')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        corrected_cell = table.rows[2].cells[1]
        corrected_para = corrected_cell.paragraphs[0]
        corrected_run = corrected_para.add_run(corr['corrected_target'])
        corrected_run.bold = True
        corrected_run.font.color.rgb = RGBColor(39, 174, 96)

        # Row 4: Explanation
        cell = table.rows[3].cells[0]
        cell.text = 'Paskaidrojums'
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, 'FFC107')

        table.rows[3].cells[1].text = corr['explanation']

        doc.add_paragraph()  # Spacer

    # Add final notes
    doc.add_page_break()
    notes_heading = doc.add_heading('Piezīmes / Notes', 1)

    notes = doc.add_paragraph()
    notes.add_run('Tulkojuma kvalitātes novērtējums:\n\n').bold = True
    notes.add_run(
        'Tulkojuma kvalitāte kopumā ir ļoti augsta (99.7%). Lielākā daļa segmentu ir iztulkoti '
        'precīzi, ievērojot terminoloģijas konsekvenci un tehniskā dokumenta prasības.\n\n'
    )

    notes.add_run('Galvenās problēmu kategorijas:\n\n').bold = True
    notes.add_run(
        '1. UI elementi bez tulkojuma (5 segmenti) - režīmu un pogu nosaukumi nav formatēti '
        'kā EN (Tulkojums), kas apgrūtina lietotāja orientēšanos starp angļu un latviešu nosaukumiem.\n\n'
        '2. Skaitļu/faktu neatbilstība (2 segmenti) - programmatūras versijas numurs un skaitļa '
        'formāta neatbilstība starp avottekstu un tulkojumu.\n\n'
        '3. Konsekvences jautājumi (1 segments) - terminu lietojuma konsekvence dažādos segmentos.\n\n'
    )

    notes.add_run('Ieteikumi:\n\n').bold = True
    notes.add_run(
        '• Ieviest konsekventu UI elementu formatēšanas praksi: objekta tips + EN nosaukums + (LV tulkojums)\n'
        '• Pārbaudīt visus skaitliskos datus un programmatūras versijas, salīdzinot ar avottekstu\n'
        '• Izveidot terminoloģijas glosāriju konsekventai terminu lietošanai\n'
        '• Veikt kvalitātes kontroli, īpaši pievēršot uzmanību UI elementiem un tehniskajiem datiem\n'
    )

    # Save document
    filename = 'LB17-11_tulkojuma_labojumi.docx'
    doc.save(filename)

    print(f"✅ DOCX fails izveidots: {filename}")
    print(f"   Faila atrašanās vieta: /home/user/TradosMatchPhrases/{filename}")
    print(f"   Iekļauti {len(corrections)} labotie segmenti")

def main():
    create_corrected_docx()

if __name__ == "__main__":
    main()
