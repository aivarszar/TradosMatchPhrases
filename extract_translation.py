#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Translation Quality Analysis Tool
Extracts and analyzes bilingual translation table from DOCX
"""
import docx
import json
import re
from collections import defaultdict

def extract_table(filename):
    """Extract bilingual table from DOCX"""
    doc = docx.Document(filename)

    if not doc.tables:
        print("No tables found in document")
        return []

    table = doc.tables[0]
    segments = []

    for idx, row in enumerate(table.rows):
        if len(row.cells) >= 2:
            source = row.cells[0].text.strip()
            target = row.cells[1].text.strip()

            if source or target:  # Skip empty rows
                segments.append({
                    'id': idx,
                    'source': source,
                    'target': target
                })

    return segments

def save_segments(segments, output_file):
    """Save segments to JSON for analysis"""
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(segments, f, ensure_ascii=False, indent=2)
    print(f"Saved {len(segments)} segments to {output_file}")

if __name__ == "__main__":
    print("Extracting translation table...")
    segments = extract_table("LB17-11 viss tulkojums.docx")
    save_segments(segments, "translation_segments.json")

    # Show some statistics
    print(f"\nTotal segments: {len(segments)}")
    print(f"Empty targets: {sum(1 for s in segments if not s['target'])}")
    print(f"Empty sources: {sum(1 for s in segments if not s['source'])}")
